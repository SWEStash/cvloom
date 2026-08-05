"""Tests for JSON Resume export."""

from __future__ import annotations

import dataclasses
from pathlib import Path

import pytest

from cvloom.export import (
    export_docx,
    export_markdown,
    export_text,
    to_json_resume,
    to_markdown,
    to_text,
)
from cvloom.models import ResolvedProfile


def _make_resolved(**overrides: object) -> ResolvedProfile:
    data = {
        "basics": {
            "headline": "Software Engineer",
            "summary": "A great engineer.",
            "links": [
                {"label": "LinkedIn", "url": "https://linkedin.com/in/janedoe"},
                {"label": "GitHub", "url": "https://github.com/janedoe"},
                {"label": "Website", "url": "https://jane.dev"},
            ],
        },
        "contact": {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "+1 555-1234",
            "location": "San Francisco, CA",
        },
        "work": [
            {
                "company": "Acme Corp",
                "title": "Senior Engineer",
                "start_date": "2021-03",
                "end_date": "Present",
                "location": "Remote",
                "highlights": ["Built scalable systems."],
            },
        ],
        "education": [
            {
                "institution": "MIT",
                "degree": "BSc",
                "field": "Computer Science",
                "start_date": "2014",
                "end_date": "2018",
                "grade": "3.9",
                "highlights": ["Dean's list."],
            },
        ],
        "skills": [
            {"category": "Languages", "items": ["Python", {"name": "Go", "level": "advanced"}]},
        ],
        "projects": [
            {
                "name": "cvloom",
                "description": "CLI CV tool.",
                "tags": ["python", "cli"],
                "url": "https://github.com/j/cvloom",
                "start_date": "2026-01",
                "highlights": ["Built it."],
            },
        ],
    }
    data.update(overrides)  # type: ignore[arg-type]
    return ResolvedProfile(
        profile={},
        data=data,
        show_sections={"work": True, "education": True, "skills": True, "projects": True},
        section_order=["skills", "work", "education", "projects"],
        template_name="cv/ats-clean",
        output_filename="cv",
    )


def test_basics_mapping():
    result = to_json_resume(_make_resolved())
    assert result["basics"]["name"] == "Jane Doe"
    assert result["basics"]["email"] == "jane@example.com"
    assert result["basics"]["label"] == "Software Engineer"
    assert result["basics"]["summary"] == "A great engineer."
    assert result["basics"]["phone"] == "+1 555-1234"


def test_location_mapping():
    result = to_json_resume(_make_resolved())
    assert result["basics"]["location"]["address"] == "San Francisco, CA"


def test_profiles_mapping():
    result = to_json_resume(_make_resolved())
    profiles = result["basics"]["profiles"]
    networks = {p["network"] for p in profiles}
    assert "LinkedIn" in networks
    assert "GitHub" in networks
    gh = next(p for p in profiles if p["network"] == "GitHub")
    assert gh["username"] == "janedoe"
    assert "github.com" in gh["url"]


def test_work_mapping():
    result = to_json_resume(_make_resolved())
    work = result["work"][0]
    assert work["name"] == "Acme Corp"
    assert work["position"] == "Senior Engineer"
    assert work["startDate"] == "2021-03"
    assert work["location"] == "Remote"
    assert "Built scalable systems." in work["highlights"]


def test_current_role_omits_end_date():
    """JSON Resume has no "Present" sentinel — a current role omits endDate."""
    work = to_json_resume(_make_resolved())["work"][0]
    assert "endDate" not in work


def test_non_iso_dates_are_omitted_not_emitted_invalid():
    resolved = _make_resolved(
        work=[{"company": "A", "title": "T", "start_date": "summer 2020", "end_date": "2021"}]
    )
    work = to_json_resume(resolved)["work"][0]
    assert "startDate" not in work
    assert work["endDate"] == "2021"


def test_education_mapping():
    result = to_json_resume(_make_resolved())
    edu = result["education"][0]
    assert edu["institution"] == "MIT"
    assert edu["studyType"] == "BSc"
    assert edu["area"] == "Computer Science"
    assert edu["score"] == "3.9"
    # JSON Resume education has no `highlights` field; `courses` is the nearest.
    assert "Dean's list." in edu["courses"]
    assert "highlights" not in edu


def test_skills_mapping():
    result = to_json_resume(_make_resolved())
    skill = result["skills"][0]
    assert skill["name"] == "Languages"
    assert "Python" in skill["keywords"]
    assert "Go" in skill["keywords"]


def test_projects_mapping():
    result = to_json_resume(_make_resolved())
    proj = result["projects"][0]
    assert proj["name"] == "cvloom"
    assert proj["description"] == "CLI CV tool."
    assert proj["url"] == "https://github.com/j/cvloom"
    assert proj["keywords"] == ["python", "cli"]
    assert "Built it." in proj["highlights"]


def test_minimal_contact():
    result = to_json_resume(
        _make_resolved(
            contact={
                "name": "Min",
                "email": "min@example.com",
            },
            basics={"headline": "Software Engineer", "summary": "A great engineer."},
        )
    )
    assert result["basics"]["name"] == "Min"
    assert "phone" not in result["basics"]
    assert "profiles" not in result["basics"]


def test_empty_sections():
    result = to_json_resume(_make_resolved(work=[], education=[], skills=[], projects=[]))
    assert "work" not in result
    assert "education" not in result
    assert "skills" not in result
    assert "projects" not in result


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------


def test_markdown_contains_name() -> None:
    md = to_markdown(_make_resolved())
    assert "# Jane Doe" in md


def test_markdown_section_headings() -> None:
    md = to_markdown(_make_resolved())
    assert "## Work Experience" in md
    assert "## Education" in md
    assert "## Skills" in md


def test_markdown_entry_content() -> None:
    md = to_markdown(_make_resolved())
    assert "Senior Engineer at Acme Corp" in md
    assert "- Built scalable systems." in md
    assert "**Languages:**" in md
    assert "Python" in md
    assert "Go" in md


def test_markdown_hidden_section_omitted() -> None:
    r = dataclasses.replace(
        _make_resolved(),
        show_sections={"work": False, "education": True, "skills": True, "projects": True},
    )
    md = to_markdown(r)
    assert "## Work Experience" not in md
    assert "## Education" in md


def test_markdown_section_order() -> None:
    r = dataclasses.replace(
        _make_resolved(),
        section_order=["education", "skills", "work", "projects"],
    )
    md = to_markdown(r)
    assert md.index("## Education") < md.index("## Skills") < md.index("## Work Experience")


def test_export_markdown_writes_file(tmp_path: Path) -> None:
    out = tmp_path / "cv.md"
    export_markdown(_make_resolved(), out)
    assert out.exists()
    assert "# Jane Doe" in out.read_text()


# ---------------------------------------------------------------------------
# Plain text export
# ---------------------------------------------------------------------------


def _text_with_every_section() -> ResolvedProfile:
    """A resolved profile carrying all nine sections, each shown."""
    extra = {
        "publications": [
            {"name": "On Scaling", "publisher": "ACM", "release_date": "2019"},
        ],
        "certifications": [
            {"name": "CKA", "issuer": "CNCF", "date": "2023-04", "type": "certification"},
        ],
        "awards": [
            {"title": "Engineer of the Year", "awarder": "Acme Corp", "date": "2022"},
        ],
        "languages": [
            {"language": "Spanish", "fluency": "Native speaker"},
        ],
    }
    every = ["skills", "work", "education", "projects"] + list(extra)
    return dataclasses.replace(
        _make_resolved(**extra),
        show_sections=dict.fromkeys(every, True),
        section_order=every,
    )


def test_text_contains_header() -> None:
    txt = to_text(_make_resolved())
    assert txt.startswith("Jane Doe\n")
    assert "Software Engineer | jane@example.com | +1 555-1234 | San Francisco, CA" in txt
    assert "https://github.com/janedoe" in txt


def test_text_section_headings() -> None:
    txt = to_text(_make_resolved())
    assert "SUMMARY\n-------" in txt
    assert "WORK EXPERIENCE\n---------------" in txt
    assert "EDUCATION\n---------" in txt
    assert "SKILLS\n------" in txt


def test_text_entry_content() -> None:
    txt = to_text(_make_resolved())
    assert "Senior Engineer at Acme Corp" in txt
    assert "2021-03 - Present | Remote" in txt
    assert "· Built scalable systems." in txt


def test_text_keeps_skill_categories() -> None:
    """The old LinkedIn export flattened every group into one line and dropped the names."""
    txt = to_text(_make_resolved())
    assert "Languages: Python, Go" in txt


def test_text_keeps_education_highlights() -> None:
    txt = to_text(_make_resolved())
    assert "· Dean's list." in txt


def test_text_renders_every_section() -> None:
    """Projects, publications, certifications, awards and languages were silently dropped."""
    txt = to_text(_text_with_every_section())
    assert "PROJECTS" in txt
    assert "cvloom" in txt and "CLI CV tool." in txt and "· Built it." in txt
    assert "PUBLICATIONS" in txt
    assert "On Scaling" in txt and "ACM · 2019" in txt
    assert "CERTIFICATIONS" in txt
    assert "· CKA | CNCF · 2023-04" in txt
    assert "AWARDS" in txt
    assert "Engineer of the Year" in txt and "Acme Corp · 2022" in txt
    assert "LANGUAGES" in txt
    assert "Spanish (Native speaker)" in txt


def test_text_honours_section_titles() -> None:
    r = dataclasses.replace(_make_resolved(), section_titles={"work": "Where I've Worked"})
    txt = to_text(r)
    assert "WHERE I'VE WORKED\n-----------------" in txt
    assert "WORK EXPERIENCE" not in txt


def test_text_hidden_section_omitted() -> None:
    r = dataclasses.replace(
        _make_resolved(),
        show_sections={"work": False, "education": True, "skills": True, "projects": True},
    )
    txt = to_text(r)
    assert "WORK EXPERIENCE" not in txt
    assert "EDUCATION" in txt


def test_text_section_order() -> None:
    r = dataclasses.replace(
        _make_resolved(),
        section_order=["education", "skills", "work", "projects"],
    )
    txt = to_text(r)
    assert txt.index("EDUCATION") < txt.index("SKILLS") < txt.index("WORK EXPERIENCE")


def test_export_text_writes_file(tmp_path: Path) -> None:
    out = tmp_path / "cv.txt"
    export_text(_make_resolved(), out)
    assert out.exists()
    assert out.read_text().startswith("Jane Doe\n")


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------


def test_export_docx_writes_file(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    out = tmp_path / "cv.docx"
    export_docx(_make_resolved(), out)
    assert out.exists()
    doc = docx.Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert any("Jane Doe" in t for t in texts)


def test_export_docx_has_section_headings(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    out = tmp_path / "cv.docx"
    export_docx(_make_resolved(), out)
    doc = docx.Document(str(out))
    styles = {p.style.name for p in doc.paragraphs}
    assert "Heading 1" in styles
    assert "List Bullet" in styles


# ── publications ─────────────────────────────────────────────────────

_PUBLICATION = {
    "name": "A model of distributed consensus under churn",
    "publisher": "Journal of Systems Research",
    "release_date": "2018",
    "identifier": "ISBN 978-0-0000-0000-1",
    "url": "https://example.com/paper",
    "summary": "A short summary of the paper.",
}


def _resolved_with_publications():
    resolved = _make_resolved(publications=[dict(_PUBLICATION)])
    resolved.show_sections["publications"] = True
    resolved.section_order.append("publications")
    return resolved


def test_publications_mapping():
    result = to_json_resume(_resolved_with_publications())
    pub = result["publications"][0]
    assert pub["name"] == _PUBLICATION["name"]
    assert pub["publisher"] == "Journal of Systems Research"
    assert pub["releaseDate"] == "2018"
    assert pub["url"] == "https://example.com/paper"


def test_publications_identifier_folded_into_summary():
    """JSON Resume has no ISBN/DOI field — export must not silently drop it."""
    result = to_json_resume(_resolved_with_publications())
    assert result["publications"][0]["summary"] == (
        "A short summary of the paper. ISBN 978-0-0000-0000-1"
    )


def test_publications_omitted_when_empty():
    assert "publications" not in to_json_resume(_make_resolved())


def test_markdown_includes_publications():
    md = to_markdown(_resolved_with_publications())
    assert "## Publications" in md
    assert _PUBLICATION["name"] in md
    assert "Journal of Systems Research · 2018 · ISBN 978-0-0000-0000-1" in md


# ── certifications ───────────────────────────────────────────────────

_CERTIFICATION = {
    "name": "AWS Certified Solutions Architect",
    "issuer": "Amazon Web Services",
    "date": "2023-04",
    "expiry_date": "2026-04",
    "identifier": "AWS-PSA-12345",
    "url": "https://example.com/verify",
}


def _resolved_with_certifications():
    resolved = _make_resolved(certifications=[dict(_CERTIFICATION)])
    resolved.show_sections["certifications"] = True
    resolved.section_order.append("certifications")
    return resolved


def test_certifications_map_to_json_resume_certificates():
    cert = to_json_resume(_resolved_with_certifications())["certificates"][0]
    assert cert["name"] == "AWS Certified Solutions Architect"
    assert cert["issuer"] == "Amazon Web Services"
    assert cert["date"] == "2023-04"
    assert cert["url"] == "https://example.com/verify"


def test_certification_extensions_ride_in_namespace():
    """expiry_date and identifier have no spec home — keep them namespaced
    rather than dropping them, so a round-trip is lossless."""
    cert = to_json_resume(_resolved_with_certifications())["certificates"][0]
    assert cert["x-cvloom-expiry_date"] == "2026-04"
    assert cert["x-cvloom-identifier"] == "AWS-PSA-12345"


def test_certificates_omitted_when_empty():
    assert "certificates" not in to_json_resume(_make_resolved())


def test_markdown_keeps_certification_extensions():
    md = to_markdown(_resolved_with_certifications())
    assert "## Certifications" in md
    assert "2023-04 - 2026-04" in md
    assert "AWS-PSA-12345" in md


# ── basics.links ─────────────────────────────────────────────────────


def _with_links(*links):
    return _make_resolved(basics={"headline": "Eng", "summary": "S", "links": list(links)})


def test_links_map_to_profiles():
    resolved = _with_links({"label": "Blog", "url": "https://example.com/blog"})
    profiles = to_json_resume(resolved)["basics"]["profiles"]
    assert {"network": "Blog", "url": "https://example.com/blog"} in profiles


def test_known_network_links_carry_a_username():
    resolved = _with_links({"label": "GitHub", "url": "https://github.com/janedoe"})
    profiles = to_json_resume(resolved)["basics"]["profiles"]
    assert profiles == [
        {
            "network": "GitHub",
            "url": "https://github.com/janedoe",
            "username": "janedoe",
        }
    ]


def test_links_deduplicate_on_normalised_url():
    """The same profile written two ways is one profile, not two."""
    resolved = _with_links(
        {"label": "GitHub", "url": "https://github.com/janedoe"},
        {"label": "GH", "url": "http://www.github.com/janedoe/"},
    )
    profiles = to_json_resume(resolved)["basics"]["profiles"]
    assert len(profiles) == 1


def test_links_falls_back_to_url_when_unlabelled():
    resolved = _with_links({"url": "https://example.com/x"})
    profiles = to_json_resume(resolved)["basics"]["profiles"]
    assert {"network": "https://example.com/x", "url": "https://example.com/x"} in profiles


def test_certification_type_survives_export():
    """`type` drives the credential/coursework heading split, so it must survive."""
    resolved = _make_resolved(
        certifications=[
            {"name": "CKA", "issuer": "CNCF", "date": "2023", "type": "certification"},
            {"name": "GenAI with LLMs", "issuer": "DeepLearning.AI", "type": "course"},
        ]
    )
    doc = to_json_resume(resolved)
    assert doc["certificates"][0]["x-cvloom-type"] == "certification"
    assert doc["certificates"][1]["x-cvloom-type"] == "course"


def test_docx_uses_one_typeface(tmp_path: Path) -> None:
    """python-docx's default theme is Calibri headings over Cambria body.

    An untouched export therefore mixes a sans and a serif that nobody chose. A CV
    should be one family, and the ATS artifact especially so.
    """
    pytest.importorskip("docx")
    import docx

    out = tmp_path / "cv.docx"
    export_docx(_make_resolved(), out)
    doc = docx.Document(str(out))
    used = {p.style.font.name for p in doc.paragraphs if p.text.strip()}
    assert used == {"Arial"}, f"DOCX mixes typefaces: {used}"


# ── section_titles overrides reach the exports ───────────────────────


def test_markdown_honours_a_profile_section_title() -> None:
    """The exports are the same document as the PDF, so they carry its wording.

    A profile that renames a heading was overridden in the HTML and silently
    reverted to the registry default in Markdown and DOCX.
    """
    resolved = dataclasses.replace(
        _make_resolved(), section_titles={"work": "Professional Experience"}
    )
    md = to_markdown(resolved)
    assert "## Professional Experience" in md
    assert "## Work Experience" not in md


def test_markdown_honours_a_summary_title_override() -> None:
    resolved = dataclasses.replace(_make_resolved(), section_titles={"summary": "About"})
    md = to_markdown(resolved)
    assert "## About" in md
    assert "## Summary" not in md


def test_markdown_headings_default_without_overrides() -> None:
    assert "## Work Experience" in to_markdown(_make_resolved())


def test_docx_honours_a_profile_section_title(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    out = tmp_path / "cv.docx"
    resolved = dataclasses.replace(
        _make_resolved(), section_titles={"work": "Professional Experience", "summary": "About"}
    )
    export_docx(resolved, out)
    texts = [p.text for p in docx.Document(str(out)).paragraphs]
    assert "Professional Experience" in texts
    assert "About" in texts
    assert "Work Experience" not in texts
