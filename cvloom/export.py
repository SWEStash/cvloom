"""Export CV data to external formats (JSON Resume, Markdown, plain text, DOCX)."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from cvloom import sections
from cvloom.links import link_username, normalize_url
from cvloom.models import ResolvedProfile
from cvloom.sections import highlight_text as _hl
from cvloom.sections import skill_name as _skill_name

# Registry headings plus `skills`, whose shape is bespoke and so is not in it.
_SECTION_HEADINGS: dict[str, str] = {
    **{s.name: s.heading for s in sections.SECTIONS},
    "skills": "Skills",
}


def _heading(resolved: ResolvedProfile, key: str, default: str) -> str:
    """Return the heading for *key*: the profile's override, else *default*.

    Mirrors the `section_title` Jinja global that does this job for the templates.
    """
    override = resolved.section_titles.get(key)
    return str(override) if override else default


def _section_heading(resolved: ResolvedProfile, section: str) -> str | None:
    """As `_heading`, for an entry in `section_order`.

    ``None`` means "not a section these exports render", which is how the callers
    skip over anything in `section_order` without a heading of its own.
    """
    default = _SECTION_HEADINGS.get(section)
    return None if default is None else _heading(resolved, section, default)


def _map_profiles(links: list[dict[str, Any]] | None) -> list[dict[str, str]]:
    """Map ``basics.links`` to JSON Resume ``basics.profiles``.

    The label stands in for ``network``. Where the handle can be recovered from the
    URL path, ``username`` is populated too; other links carry a URL only.
    """
    profiles: list[dict[str, str]] = []
    seen: set[str] = set()
    for link in links or []:
        url = str(link.get("url", ""))
        label = str(link.get("label", ""))
        key = normalize_url(url)
        if not url or key in seen:
            continue
        profile = {"network": label or url, "url": url}
        username = link_username(url)
        if username:
            profile["username"] = username
        profiles.append(profile)
        seen.add(key)
    return profiles


def _map_location(contact: dict[str, Any]) -> dict[str, str]:
    """Map location string to JSON Resume location object."""
    loc = contact.get("location", "")
    return {"address": loc} if loc else {}


# ── JSON Resume field mapping ────────────────────────────────────────
#
# Each cvloom section maps to a JSON Resume array by renaming fields, so a new
# section or field is a table edit.

# JSON Resume requires ISO 8601 with flexible granularity.
_ISO_DATE_RE = re.compile(r"^\d{4}(-\d{2}(-\d{2})?)?$")

# Tags have no JSON Resume home outside projects. A namespaced extension keeps
# them through a round-trip; conforming consumers ignore unknown x- keys.
TAGS_EXTENSION_KEY = "x-cvloom-tags"


@dataclass(frozen=True)
class _Field:
    """One cvloom key → JSON Resume key mapping."""

    src: str
    dest: str
    kind: str = "text"  # text | date | list | highlights


def _iso_date(value: Any) -> str | None:
    """Return *value* if it is a valid JSON Resume date, else None.

    cvloom allows free text such as ``"Present"``; JSON Resume expresses a current
    role by omitting endDate, so anything non-conforming is dropped.
    """
    text = str(value or "")
    return text if _ISO_DATE_RE.match(text) else None


def _map_entry(entry: dict[str, Any], fields: tuple[_Field, ...]) -> dict[str, Any]:
    """Apply *fields* to one entry, omitting anything empty or non-conforming."""
    item: dict[str, Any] = {}
    for f in fields:
        value = entry.get(f.src)
        if f.kind == "date":
            iso = _iso_date(value)
            if iso:
                item[f.dest] = iso
        elif f.kind == "highlights":
            if value:
                item[f.dest] = [_hl(h) for h in value]
        elif f.kind == "list":
            if value:
                item[f.dest] = list(value)
        elif value:
            item[f.dest] = value
    if entry.get("tags") and not any(f.src == "tags" for f in fields):
        item[TAGS_EXTENSION_KEY] = list(entry["tags"])
    return item


def _map_entries(entries: list[dict[str, Any]], fields: tuple[_Field, ...]) -> list[dict[str, Any]]:
    return [_map_entry(entry, fields) for entry in entries]


_WORK_FIELDS = (
    _Field("company", "name"),
    _Field("title", "position"),
    _Field("location", "location"),
    _Field("start_date", "startDate", "date"),
    _Field("end_date", "endDate", "date"),
    _Field("highlights", "highlights", "highlights"),
)

_EDUCATION_FIELDS = (
    _Field("institution", "institution"),
    _Field("degree", "studyType"),
    _Field("field", "area"),
    _Field("start_date", "startDate", "date"),
    _Field("end_date", "endDate", "date"),
    _Field("grade", "score"),
    # JSON Resume education has no `highlights`; `courses` is the nearest field.
    _Field("highlights", "courses", "highlights"),
)

_PROJECT_FIELDS = (
    _Field("name", "name"),
    _Field("description", "description"),
    _Field("url", "url"),
    _Field("start_date", "startDate", "date"),
    _Field("end_date", "endDate", "date"),
    _Field("tags", "keywords", "list"),
    _Field("highlights", "highlights", "highlights"),
)

_PUBLICATION_FIELDS = (
    _Field("name", "name"),
    _Field("publisher", "publisher"),
    _Field("release_date", "releaseDate", "date"),
    _Field("url", "url"),
)

_AWARD_FIELDS = (
    _Field("title", "title"),
    _Field("awarder", "awarder"),
    _Field("date", "date", "date"),
    _Field("summary", "summary"),
)

_LANGUAGE_FIELDS = (
    _Field("language", "language"),
    _Field("fluency", "fluency"),
)

_CERTIFICATION_FIELDS = (
    _Field("name", "name"),
    _Field("issuer", "issuer"),
    _Field("date", "date", "date"),
    _Field("url", "url"),
)


def _map_skills(groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Map skill groups to JSON Resume skills array.

    Per-item proficiency has no JSON Resume home, so it rides in the extension
    namespace.
    """
    result: list[dict[str, Any]] = []
    for group in groups:
        items = group.get("items", [])
        entry: dict[str, Any] = {
            "name": group.get("category", ""),
            "keywords": [_skill_name(item) for item in items],
        }
        levels = {
            _skill_name(i): i["level"] for i in items if isinstance(i, dict) and i.get("level")
        }
        if levels:
            entry["x-cvloom-levels"] = levels
        result.append(entry)
    return result


def _map_publications(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Map publication entries to the JSON Resume publications array.

    ``identifier`` (ISBN/DOI) has no counterpart, so it is folded into ``summary``.
    """
    items = _map_entries(entries, _PUBLICATION_FIELDS)
    for item, entry in zip(items, entries, strict=True):
        parts = [p for p in (entry.get("summary"), entry.get("identifier")) if p]
        if parts:
            item["summary"] = " ".join(parts)
    return items


def _map_certifications(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Map certification entries to the JSON Resume certificates array.

    JSON Resume's certificate object is only {name, date, issuer, url}, so
    ``expiry_date``, ``identifier`` and ``type`` ride in the extension namespace.
    """
    items = _map_entries(entries, _CERTIFICATION_FIELDS)
    for item, entry in zip(items, entries, strict=True):
        for key in ("expiry_date", "identifier", "type"):
            if entry.get(key):
                item[f"x-cvloom-{key}"] = entry[key]
    return items


def _meta_line(*parts: Any) -> str:
    """Join the present values with a separator: `ACM SIGPLAN · 2019`."""
    return " · ".join(str(p) for p in parts if p)


def _language_label(entry: dict[str, Any]) -> str:
    """`Spanish (Native speaker)` — fluency in parentheses when present."""
    name = str(entry.get("language", ""))
    fluency = entry.get("fluency")
    return f"{name} ({fluency})" if fluency else name


def _certification_dates(entry: dict[str, Any]) -> str:
    """`2023-04` or `2023-04 - 2026-04` when the credential expires."""
    date = entry.get("date", "")
    expiry = entry.get("expiry_date")
    return f"{date} - {expiry}" if date and expiry else str(date)


def _certification_meta(entry: dict[str, Any]) -> str:
    return _meta_line(entry.get("issuer"), _certification_dates(entry), entry.get("identifier"))


def _certification_line(entry: dict[str, Any]) -> str:
    """One-line rendering of a certification: name, issuer, dates, ID."""
    meta = _certification_meta(entry)
    return f"**{entry.get('name', '')}**{f' | {meta}' if meta else ''}"


def _header_parts(resolved: ResolvedProfile, *, bold_headline: bool = False) -> list[str]:
    """The contact line under the name: headline, reachability, then every link URL."""
    contact = resolved.data.get("contact", {})
    basics = resolved.data.get("basics", {})

    parts: list[str] = []
    if basics.get("headline"):
        headline = str(basics["headline"])
        parts.append(f"**{headline}**" if bold_headline else headline)
    for field in ["email", "phone", "location"]:
        if contact.get(field):
            parts.append(str(contact[field]))
    parts += [str(link["url"]) for link in basics.get("links", []) if link.get("url")]
    return parts


def to_json_resume(resolved: ResolvedProfile) -> dict[str, Any]:
    """Convert resolved cvloom data to JSON Resume schema."""
    data = resolved.data
    contact = data.get("contact", {})
    basics_data = data.get("basics", {})

    # Only emit fields that carry a value: a --public build strips email, and
    # an empty string fails JSON Resume's `email` format constraint.
    basics: dict[str, Any] = {}
    for key, value in (
        ("name", contact.get("name")),
        ("label", basics_data.get("headline")),
        ("email", contact.get("email")),
        ("phone", contact.get("phone")),
        ("summary", basics_data.get("summary")),
    ):
        if value:
            basics[key] = value
    result: dict[str, Any] = {"basics": basics}

    location = _map_location(contact)
    if location:
        result["basics"]["location"] = location

    profiles = _map_profiles(basics_data.get("links"))
    if profiles:
        result["basics"]["profiles"] = profiles

    work = data.get("work", [])
    if work:
        result["work"] = _map_entries(work, _WORK_FIELDS)

    education = data.get("education", [])
    if education:
        result["education"] = _map_entries(education, _EDUCATION_FIELDS)

    skills = data.get("skills", [])
    if skills:
        result["skills"] = _map_skills(skills)

    projects = data.get("projects", [])
    if projects:
        result["projects"] = _map_entries(projects, _PROJECT_FIELDS)

    publications = data.get("publications", [])
    if publications:
        result["publications"] = _map_publications(publications)

    certifications = data.get("certifications", [])
    if certifications:
        result["certificates"] = _map_certifications(certifications)

    awards = data.get("awards", [])
    if awards:
        result["awards"] = _map_entries(awards, _AWARD_FIELDS)

    languages = data.get("languages", [])
    if languages:
        result["languages"] = _map_entries(languages, _LANGUAGE_FIELDS)

    return result


def export_json_resume(resolved: ResolvedProfile, output_path: Path) -> None:
    """Write JSON Resume file."""
    resume = to_json_resume(resolved)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(resume, indent=2, ensure_ascii=False) + "\n")


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def to_markdown(resolved: ResolvedProfile) -> str:
    """Return the CV formatted as a Markdown string."""
    data = resolved.data
    contact = data.get("contact", {})
    basics = data.get("basics", {})

    lines: list[str] = [f"# {contact.get('name', '')}", ""]

    parts = _header_parts(resolved, bold_headline=True)
    if parts:
        lines += [" | ".join(parts), ""]

    if basics.get("summary"):
        lines += [
            "---",
            "",
            f"## {_heading(resolved, 'summary', 'Summary')}",
            "",
            str(basics["summary"]),
            "",
        ]

    for section in resolved.section_order:
        if not resolved.show_sections.get(section, False):
            continue
        heading = _section_heading(resolved, section)
        if heading is None:
            continue

        lines += ["---", "", f"## {heading}", ""]

        if section == "skills":
            for group in data.get("skills", []):
                cat = group.get("category", "")
                items = [_skill_name(i) for i in group.get("items", [])]
                lines.append(f"**{cat}:** {', '.join(items)}")
            lines.append("")

        elif section == "work":
            for entry in data.get("work", []):
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                location = entry.get("location", "")
                date_str = f"{start} - {end}" if end else start
                meta = f"{date_str} | {location}" if location else date_str
                lines.append(f"### {entry.get('title', '')} at {entry.get('company', '')}")
                lines += [f"*{meta}*", ""]
                for h in entry.get("highlights", []):
                    lines.append(f"- {_hl(h)}")
                lines.append("")

        elif section == "education":
            for entry in data.get("education", []):
                degree = entry.get("degree", "")
                field = entry.get("field", "")
                institution = entry.get("institution", "")
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                location = entry.get("location", "")
                degree_str = f"{degree} in {field}" if field else degree
                date_str = f"{start} - {end}" if end else start
                meta = f"{date_str} | {location}" if location else date_str
                lines.append(f"### {degree_str} | {institution}")
                lines += [f"*{meta}*", ""]
                for h in entry.get("highlights", []):
                    lines.append(f"- {_hl(h)}")
                lines.append("")

        elif section == "projects":
            for entry in data.get("projects", []):
                name = entry.get("name", "")
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                description = entry.get("description", "")
                lines.append(f"### {name}")
                if start:
                    date_str = f"{start} - {end}" if end else start
                    lines += [f"*{date_str}*", ""]
                else:
                    lines.append("")
                if description:
                    lines += [description, ""]
                for h in entry.get("highlights", []):
                    lines.append(f"- {_hl(h)}")
                lines.append("")

        elif section == "publications":
            for entry in data.get("publications", []):
                meta = _meta_line(
                    entry.get("publisher"), entry.get("release_date"), entry.get("identifier")
                )
                lines.append(f"### {entry.get('name', '')}")
                lines += [f"*{meta}*", ""] if meta else [""]
                if entry.get("summary"):
                    lines += [str(entry["summary"]), ""]

        elif section == "certifications":
            for entry in data.get("certifications", []):
                lines.append(f"- {_certification_line(entry)}")
            lines.append("")

        elif section == "awards":
            for entry in data.get("awards", []):
                meta = _meta_line(entry.get("awarder"), entry.get("date"))
                lines.append(f"### {entry.get('title', '')}")
                lines += [f"*{meta}*", ""] if meta else [""]
                if entry.get("summary"):
                    lines += [str(entry["summary"]), ""]

        elif section == "languages":
            langs = [_language_label(e) for e in data.get("languages", [])]
            lines += [" · ".join(langs), ""]

    return "\n".join(lines).rstrip() + "\n"


def export_markdown(resolved: ResolvedProfile, output_path: Path) -> None:
    """Write Markdown CV to output_path."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(to_markdown(resolved), encoding="utf-8")


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def _text_block(heading: str, *body: str) -> str:
    """A heading in caps over a rule of its own width, then the body lines."""
    banner = heading.upper()
    return "\n".join([banner, "-" * len(banner), *body]).rstrip()


def to_text(resolved: ResolvedProfile) -> str:
    """Return the CV as plain text, carrying the same content as the Markdown export."""
    data = resolved.data
    contact = data.get("contact", {})
    basics = data.get("basics", {})

    blocks: list[str] = []

    header = [contact.get("name", "")]
    parts = _header_parts(resolved)
    if parts:
        header.append(" | ".join(parts))
    blocks.append("\n".join(header).rstrip())

    if basics.get("summary"):
        blocks.append(_text_block(_heading(resolved, "summary", "Summary"), str(basics["summary"])))

    for section in resolved.section_order:
        if not resolved.show_sections.get(section, False):
            continue
        heading = _section_heading(resolved, section)
        if heading is None:
            continue

        body: list[str] = []

        if section == "skills":
            for group in data.get("skills", []):
                items = [_skill_name(i) for i in group.get("items", [])]
                body.append(f"{group.get('category', '')}: {', '.join(items)}")

        elif section == "work":
            for entry in data.get("work", []):
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                location = entry.get("location", "")
                date_str = f"{start} - {end}" if end else start
                meta = f"{date_str} | {location}" if location else date_str
                body.append(f"{entry.get('title', '')} at {entry.get('company', '')}")
                body += [meta, ""]
                body += [f"· {_hl(h)}" for h in entry.get("highlights", [])]
                body.append("")

        elif section == "education":
            for entry in data.get("education", []):
                degree = entry.get("degree", "")
                field = entry.get("field", "")
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                location = entry.get("location", "")
                degree_str = f"{degree} in {field}" if field else degree
                date_str = f"{start} - {end}" if end else start
                meta = f"{date_str} | {location}" if location else date_str
                body.append(f"{degree_str} | {entry.get('institution', '')}")
                body += [meta, ""]
                body += [f"· {_hl(h)}" for h in entry.get("highlights", [])]
                body.append("")

        elif section == "projects":
            for entry in data.get("projects", []):
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                body.append(str(entry.get("name", "")))
                if start:
                    body.append(f"{start} - {end}" if end else start)
                body.append("")
                if entry.get("description"):
                    body += [str(entry["description"]), ""]
                body += [f"· {_hl(h)}" for h in entry.get("highlights", [])]
                body.append("")

        elif section == "publications":
            for entry in data.get("publications", []):
                meta = _meta_line(
                    entry.get("publisher"), entry.get("release_date"), entry.get("identifier")
                )
                body.append(str(entry.get("name", "")))
                body += [meta, ""] if meta else [""]
                if entry.get("summary"):
                    body += [str(entry["summary"]), ""]

        elif section == "certifications":
            for entry in data.get("certifications", []):
                meta = _certification_meta(entry)
                name = entry.get("name", "")
                body.append(f"· {name} | {meta}" if meta else f"· {name}")

        elif section == "awards":
            for entry in data.get("awards", []):
                meta = _meta_line(entry.get("awarder"), entry.get("date"))
                body.append(str(entry.get("title", "")))
                body += [meta, ""] if meta else [""]
                if entry.get("summary"):
                    body += [str(entry["summary"]), ""]

        elif section == "languages":
            body.append(" · ".join(_language_label(e) for e in data.get("languages", [])))

        blocks.append(_text_block(heading, *body))

    return "\n\n\n".join(blocks) + "\n"


def export_text(resolved: ResolvedProfile, output_path: Path) -> None:
    """Write the plain-text CV to output_path."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(to_text(resolved), encoding="utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


# One typeface across the document: python-docx's default template otherwise
# inherits the Office theme, giving sans headings and serif body text.
_DOCX_FONT = "Arial"


def _pin_docx_font(doc: object) -> None:
    """Force every style used by the export onto one typeface.

    Setting `style.font.name` alone only writes `w:ascii`; Word resolves the other
    scripts through the theme, so the remaining `rFonts` attributes must be set on
    the element directly and the theme-linked ones cleared.
    """
    from docx.oxml.ns import qn

    for name in (
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Body Text",
        "List Bullet",
    ):
        try:
            style = doc.styles[name]  # type: ignore[attr-defined]
        except KeyError:
            continue
        style.font.name = _DOCX_FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rfonts.set(qn(f"w:{attr}"), _DOCX_FONT)
        for themed in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
            if rfonts.get(qn(f"w:{themed}")) is not None:
                del rfonts.attrib[qn(f"w:{themed}")]


def export_docx(resolved: ResolvedProfile, output_path: Path) -> None:
    """Write a .docx file using python-docx (optional dependency)."""
    try:
        import docx
    except ImportError:
        raise SystemExit(
            "python-docx is not installed. Install it with: uv pip install python-docx"
        )

    data = resolved.data
    contact = data.get("contact", {})
    basics = data.get("basics", {})

    doc = docx.Document()
    _pin_docx_font(doc)
    doc.add_paragraph(contact.get("name", ""), style="Title")

    parts = _header_parts(resolved)
    if parts:
        doc.add_paragraph(" | ".join(parts), style="Subtitle")

    if basics.get("summary"):
        doc.add_heading(_heading(resolved, "summary", "Summary"), level=1)
        doc.add_paragraph(str(basics["summary"]), style="Body Text")

    for section in resolved.section_order:
        if not resolved.show_sections.get(section, False):
            continue
        heading = _section_heading(resolved, section)
        if heading is None:
            continue

        doc.add_heading(heading, level=1)

        if section == "work":
            for entry in data.get("work", []):
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                location = entry.get("location", "")
                date_str = f"{start} - {end}" if end else start
                meta = f"{date_str} | {location}" if location else date_str
                doc.add_heading(f"{entry.get('title', '')} at {entry.get('company', '')}", level=2)
                if meta:
                    p = doc.add_paragraph(meta, style="Body Text")
                    p.runs[0].italic = True
                for h in entry.get("highlights", []):
                    doc.add_paragraph(_hl(h), style="List Bullet")

        elif section == "education":
            for entry in data.get("education", []):
                degree = entry.get("degree", "")
                field = entry.get("field", "")
                institution = entry.get("institution", "")
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                degree_str = f"{degree} in {field}" if field else degree
                doc.add_heading(f"{degree_str} | {institution}", level=2)
                if start:
                    p = doc.add_paragraph(f"{start} - {end}" if end else start, style="Body Text")
                    p.runs[0].italic = True
                for h in entry.get("highlights", []):
                    doc.add_paragraph(_hl(h), style="List Bullet")

        elif section == "skills":
            for group in data.get("skills", []):
                cat = group.get("category", "")
                items = [_skill_name(i) for i in group.get("items", [])]
                doc.add_paragraph(f"{cat}: {', '.join(items)}", style="Body Text")

        elif section == "projects":
            for entry in data.get("projects", []):
                name = entry.get("name", "")
                start = entry.get("start_date", "")
                end = entry.get("end_date", "")
                description = entry.get("description", "")
                doc.add_heading(name, level=2)
                if start:
                    p = doc.add_paragraph(f"{start} - {end}" if end else start, style="Body Text")
                    p.runs[0].italic = True
                if description:
                    doc.add_paragraph(description, style="Body Text")
                for h in entry.get("highlights", []):
                    doc.add_paragraph(_hl(h), style="List Bullet")

        elif section == "publications":
            for entry in data.get("publications", []):
                meta = _meta_line(
                    entry.get("publisher"), entry.get("release_date"), entry.get("identifier")
                )
                doc.add_heading(entry.get("name", ""), level=2)
                if meta:
                    p = doc.add_paragraph(meta, style="Body Text")
                    p.runs[0].italic = True
                if entry.get("summary"):
                    doc.add_paragraph(str(entry["summary"]), style="Body Text")

        elif section == "certifications":
            for entry in data.get("certifications", []):
                meta = _certification_meta(entry)
                text = entry.get("name", "")
                doc.add_paragraph(f"{text} | {meta}" if meta else text, style="List Bullet")

        elif section == "awards":
            for entry in data.get("awards", []):
                meta = _meta_line(entry.get("awarder"), entry.get("date"))
                doc.add_heading(entry.get("title", ""), level=2)
                if meta:
                    para = doc.add_paragraph(meta, style="Body Text")
                    para.runs[0].italic = True
                if entry.get("summary"):
                    doc.add_paragraph(str(entry["summary"]), style="Body Text")

        elif section == "languages":
            doc.add_paragraph(
                " · ".join(_language_label(e) for e in data.get("languages", [])),
                style="Body Text",
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
